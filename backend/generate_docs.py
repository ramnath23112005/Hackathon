from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime, os

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)

title = doc.add_heading("InternetOS", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph("AI-Powered Autonomous Internet Research Agent")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(14)
sub.runs[0].font.color.rgb = RGBColor(100, 100, 100)
doc.add_paragraph()

now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
doc.add_paragraph(f"Documentation generated: {now}")
doc.add_paragraph("Phase: 1 (Foundation) + Phase 2 (Wire Integration)")
doc.add_paragraph()

doc.add_heading("1. Project Overview", level=1)
doc.add_paragraph(
    "InternetOS is an AI agent platform that treats the internet like a structured database. "
    "Instead of manually browsing websites, scraping pages, or handling authentication flows, "
    "InternetOS uses Anakin Wire as the execution and query layer to search, analyze, monitor, "
    "and synthesize information into actionable insights."
)
doc.add_paragraph(
    "The system combines live internet data, structured APIs, autonomous AI reasoning, "
    "and multi-source synthesis into one intelligent workflow."
)

doc.add_heading("2. System Architecture", level=1)
doc.add_paragraph("The platform follows a clean layered architecture:")

items = [
    ("Frontend (Next.js)", "React-based chat UI with Tailwind CSS. Communicates with backend via HTTP POST."),
    ("Backend (FastAPI)", "Python async server with Routes, Services, Core modular structure."),
    ("AI/Reasoning Layer", "Planned for Phase 3. Structured data prepared for LLM ingestion."),
    ("Data Layer (Wire)", "Anakin Wire SDK provides structured data from Reddit, Google News, and web search."),
]
for t, d in items:
    p = doc.add_paragraph()
    r = p.add_run(f"{t}: ")
    r.bold = True
    p.add_run(d)

doc.add_paragraph()
p = doc.add_paragraph("Data Flow:")
p.runs[0].bold = True
doc.add_paragraph(
    "User Prompt -> FastAPI /chat endpoint -> Wire Service -> Wire Client -> "
    "Anakin SDK -> Structured JSON -> Normalized Response -> Frontend Display",
    style="List Bullet",
)

doc.add_heading("3. Project Folder Structure", level=1)
doc.add_paragraph("The repository is organized for production scalability:")

structure = """internetos/
  frontend/
    app/
      globals.css
      layout.tsx
      page.tsx
    components/
      ChatUI.tsx
    public/
    .env.local.example
    next.config.js
    package.json
    postcss.config.js
    tailwind.config.ts
    tsconfig.json
    vercel.json
  backend/
    core/
      config.py
      cors.py
      logger.py
      wire_client.py
    models/
      schemas.py
    routes/
      chat.py
    services/
      chat_service.py
      wire_service.py
    main.py
    Procfile
    requirements.txt
    .env.example
  .gitignore
  README.md"""

p = doc.add_paragraph()
r = p.add_run(structure)
r.font.name = "Consolas"
r.font.size = Pt(8)

doc.add_heading("4. Backend Implementation", level=1)

doc.add_heading("4.1 main.py - Application Entry Point", level=2)
doc.add_paragraph(
    "Creates the FastAPI application, configures CORS middleware from settings, "
    "includes the chat router, and exposes a /health endpoint."
)

doc.add_heading("4.2 core/config.py - Settings", level=2)
doc.add_paragraph(
    "Uses pydantic-settings to load configuration from environment variables. "
    "Supports JSON-encoded CORS origins, Wire API credentials, and action IDs."
)

doc.add_heading("4.3 core/cors.py - CORS Middleware", level=2)
doc.add_paragraph(
    "Configures FastAPI CORSMiddleware with allowed origins from settings. "
    "Supports credentials, all methods, and all headers."
)

doc.add_heading("4.4 core/logger.py - Debug Logging", level=2)
doc.add_paragraph("Provides structured logging with timestamps to stdout for pipeline debugging.")

doc.add_heading("4.5 core/wire_client.py - Anakin Wire Client", level=2)
doc.add_paragraph(
    "Wraps the official Anakin SDK. Provides async search_all() that runs "
    "Wire actions and SDK search in parallel via asyncio.to_thread(). "
    "Handles missing API key, action failures, and timeouts gracefully."
)
doc.add_paragraph("_get_client() - Creates authenticated Anakin SDK client", style="List Bullet")
doc.add_paragraph("_run_wire_action() - Executes a Wire action and returns structured result", style="List Bullet")
doc.add_paragraph("_run_web_search() - Uses SDK built-in search() for general web results", style="List Bullet")
doc.add_paragraph("search_all() - Parallel execution of all configured sources", style="List Bullet")

doc.add_heading("4.6 models/schemas.py - Data Models", level=2)
doc.add_paragraph("Pydantic v2 models for request/response validation:")
doc.add_paragraph("ChatRequest - { message: str }", style="List Bullet")
doc.add_paragraph("SourceSummary - { source, error, count }", style="List Bullet")
doc.add_paragraph("NormalizedData - { summary_context, key_entities, trends, sources, results }", style="List Bullet")
doc.add_paragraph("ChatResponse - { query, sources, data, raw_wire_response (debug) }", style="List Bullet")

doc.add_heading("4.7 routes/chat.py - API Route", level=2)
doc.add_paragraph(
    "Single POST /chat endpoint. Accepts ChatRequest, passes to chat_service, "
    "returns ChatResponse. Debug flag from settings controls raw Wire data inclusion."
)

doc.add_heading("4.8 services/chat_service.py - Orchestration", level=2)
doc.add_paragraph(
    "Coordinates the Wire pipeline: calls query_internet(), constructs "
    "NormalizedData from the response, and wraps it in ChatResponse."
)

doc.add_heading("4.9 services/wire_service.py - Data Normalization", level=2)
doc.add_paragraph("Transforms raw Wire API output into clean, structured format:")
doc.add_paragraph("Entity extraction - Captures @mentions, #hashtags, and URLs", style="List Bullet")
doc.add_paragraph("Source tracking - Maps action IDs to human-readable names", style="List Bullet")
doc.add_paragraph("Summary generation - Concatenates top 5 result texts", style="List Bullet")
doc.add_paragraph("Trend detection - Extracts hashtags as trend indicators", style="List Bullet")
doc.add_paragraph("Result limiting - Caps at 30 results, 500 chars per entry", style="List Bullet")

doc.add_heading("5. Frontend Implementation", level=1)
doc.add_heading("5.1 Tech Stack", level=2)
doc.add_paragraph("Next.js 14 (App Router) + TypeScript + Tailwind CSS", style="List Bullet")

doc.add_heading("5.2 ChatUI Component", level=2)
doc.add_paragraph("Client-side React component with the following features:")
doc.add_paragraph("Message history with user/assistant roles", style="List Bullet")
doc.add_paragraph("API_URL configurable via NEXT_PUBLIC_API_URL env var", style="List Bullet")
doc.add_paragraph("Loading indicator during backend processing", style="List Bullet")
doc.add_paragraph("Error handling with user-friendly messages", style="List Bullet")
doc.add_paragraph("Auto-scroll to latest message", style="List Bullet")
doc.add_paragraph("Structured data display: sources, summary, entities, results", style="List Bullet")

doc.add_heading("6. API Endpoints", level=1)
table = doc.add_table(rows=3, cols=3)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "Method"
hdr[1].text = "Path"
hdr[2].text = "Description"
row1 = table.rows[1].cells
row1[0].text = "GET"
row1[1].text = "/health"
row1[2].text = "Health check"
row2 = table.rows[2].cells
row2[0].text = "POST"
row2[1].text = "/chat"
row2[2].text = "Research query -> structured response"

doc.add_heading("7. Deployment", level=1)
doc.add_heading("7.1 Backend (Render)", level=2)
doc.add_paragraph("Platform: Render Web Service (Python 3)", style="List Bullet")
doc.add_paragraph("Root Directory: backend/", style="List Bullet")
doc.add_paragraph("Build Command: pip install -r requirements.txt", style="List Bullet")
doc.add_paragraph("Start Command: uvicorn main:app --host 0.0.0.0 --port $PORT", style="List Bullet")
doc.add_paragraph("Live URL: https://hackathon-9a3r.onrender.com", style="List Bullet")

doc.add_heading("7.2 Frontend (Vercel)", level=2)
doc.add_paragraph("Platform: Vercel (Next.js)", style="List Bullet")
doc.add_paragraph("Root Directory: frontend/", style="List Bullet")
doc.add_paragraph("Env Var: NEXT_PUBLIC_API_URL = backend URL", style="List Bullet")
doc.add_paragraph("Live URL: https://hackathon-navy-eta.vercel.app", style="List Bullet")

doc.add_heading("8. Environment Variables", level=1)
env_table = doc.add_table(rows=len(envs)+1, cols=2)
env_table.style = "Light Grid Accent 1"
env_hdr = env_table.rows[0].cells
env_hdr[0].text = "Variable"
env_hdr[1].text = "Purpose"
envs = [
    ("APP_NAME", "Application display name"),
    ("APP_VERSION", "API version string"),
    ("DEBUG", "Enable debug mode (true/false)"),
    ("CORS_ORIGINS", "JSON array of allowed origins"),
    ("WIRE_API_KEY", "Anakin API key"),
    ("WIRE_BASE_URL", "Anakin API base URL"),
    ("WIRE_ACTION_REDDIT", "Reddit Wire action ID (rt_search)"),
    ("WIRE_ACTION_NEWS", "Google News Wire action ID (gn_search)"),
]
for i, (k, v) in enumerate(envs, 1):
    env_table.rows[i].cells[0].text = k
    env_table.rows[i].cells[1].text = v

doc.add_heading("9. Local Development Setup", level=1)
doc.add_paragraph("Backend:", style="List Bullet")
doc.add_paragraph("cd backend", style="List Bullet 2")
doc.add_paragraph("pip install -r requirements.txt", style="List Bullet 2")
doc.add_paragraph("uvicorn main:app --reload --port 8000", style="List Bullet 2")
doc.add_paragraph("Frontend:", style="List Bullet")
doc.add_paragraph("cd frontend", style="List Bullet 2")
doc.add_paragraph("npm install", style="List Bullet 2")
doc.add_paragraph("npm run dev", style="List Bullet 2")
doc.add_paragraph("Open http://localhost:3000", style="List Bullet 2")

doc.add_heading("10. Common Issues & Fixes", level=1)
issues = [
    ("CORS Error", "Ensure CORS_ORIGINS includes the frontend URL on Render"),
    ("Port Conflict", "Kill existing uvicorn: taskkill /F /IM uvicorn.exe"),
    ("Wire API Failure", "Check WIRE_API_KEY is set. SDK returns empty results gracefully."),
    ("No Results from Action", "Verify action ID. Some actions may have auth issues on Anakin side."),
    ("Frontend Stale", "Set NEXT_PUBLIC_API_URL in Vercel env vars and redeploy."),
    ("Build Fails", "Ensure root directory is backend/ and build command is pip install."),
]
for t, d in issues:
    p = doc.add_paragraph()
    r = p.add_run(f"{t}: ")
    r.bold = True
    p.add_run(d)

doc.add_heading("11. Git Commit History", level=1)
commits = [
    ("Phase 1", "25 files - Initial scaffold: Next.js + FastAPI + chat UI"),
    ("Phase 1", "Procfile for Render deployment"),
    ("Phase 1", "vercel.json + env var for API URL"),
    ("Phase 1", "CORS origins JSON parsing fix"),
    ("Phase 2", "Wire API integration: structured data pipeline (10 files)"),
    ("Phase 2", "Use official Anakin SDK for Wire integration"),
    ("Phase 2", "Remove X/Twitter, use SDK search for web"),
    ("Phase 2", "Fix source names, remove X from loading text"),
]
for phase, desc in commits:
    p = doc.add_paragraph()
    r = p.add_run(f"[{phase}] ")
    r.bold = True
    p.add_run(desc)

output_path = "C:/Users/ramna/Desktop/Hackathon/internetos/docs/InternetOS_Documentation.docx"
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f"Document saved to: {output_path}")
